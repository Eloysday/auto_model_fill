#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
JSON 还原为 Word 文档脚本（增强版）
将填报完成的 template_full.json 还原为格式保持的 .docx 文件

支持:
  - 段落 / 表格 / 图片（含图表输出）
  - 页面设置、页边距
  - 样式还原（字体、字号、颜色、粗斜体、下划线等）
  - 单元格底色、边框、边距
  - 项目编号（numbering）
  - 图片嵌入（从本地路径）
  python restore_report.py template.docx fill_test.json --out report.docx
"""

import os, json, copy
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement


# ═══════════════════════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════════════════════

def hex_to_rgb(hex_color):
    """十六进制颜色 → RGB tuple，失败或 'auto' 返回 None"""
    if not hex_color or hex_color.lower() == 'auto':
        return None
    hex_color = hex_color.lstrip('#')
    if len(hex_color) == 6:
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    return None


def set_run_font(run, fmt, defaults=None):
    """设置单个 <w:r> 的字体格式"""
    if not fmt:
        return

    rPr = run._element.get_or_add_rPr()

    # ── 字体大小 (sz_hp = 半磅, 例如 22 → 11pt) ──
    sz_hp = fmt.get('sz_hp')
    if sz_hp:
        run.font.size = Pt(sz_hp / 2.0)
    szCs_hp = fmt.get('szCs_hp')
    if szCs_hp:
        el = rPr.find(qn('w:szCs'))
        if el is None:
            el = OxmlElement('w:szCs')
            rPr.append(el)
        el.set(qn('w:val'), str(int(szCs_hp)))

    # ── 颜色 ──
    color = fmt.get('color')
    if color and color.lower() != 'auto':
        rgb = hex_to_rgb(color)
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)

    # ── 粗体 ──
    if fmt.get('bold') is not None:
        run.font.bold = fmt['bold']
    if fmt.get('bold_cjk') is not None:
        el = rPr.find(qn('w:bCs'))
        if el is None:
            el = OxmlElement('w:bCs')
            rPr.append(el)
        if not fmt['bold_cjk']:
            el.set(qn('w:val'), '0')

    # ── 斜体 ──
    if fmt.get('italic') is not None:
        run.font.italic = fmt['italic']
    if fmt.get('italic_cjk') is not None:
        el = rPr.find(qn('w:iCs'))
        if el is None:
            el = OxmlElement('w:iCs')
            rPr.append(el)
        if not fmt['italic_cjk']:
            el.set(qn('w:val'), '0')

    # ── 下划线 ──
    underline = fmt.get('underline')
    if underline:
        run.font.underline = True

    # ── 字体名称 ──
    font = fmt.get('font', {})
    if font:
        name = font.get('ascii') or font.get('hAnsi') or font.get('eastAsia')
        if name:
            run.font.name = name
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            # 先清除主题字体引用
            for theme_attr in ('w:asciiTheme', 'w:eastAsiaTheme', 'w:hAnsiTheme', 'w:cstheme'):
                if rFonts.get(qn(theme_attr)):
                    del rFonts.attrib[qn(theme_attr)]
            rFonts.set(qn('w:eastAsia'), font.get('eastAsia', name))
            if font.get('ascii'):
                rFonts.set(qn('w:ascii'), font['ascii'])
            if font.get('hAnsi'):
                rFonts.set(qn('w:hAnsi'), font['hAnsi'])
            if font.get('cs'):
                rFonts.set(qn('w:cs'), font['cs'])

    # ── 上/下标 ──
    va = fmt.get('vertAlign')
    if va:
        run.font.superscript = (va == 'superscript')
        run.font.subscript = (va == 'subscript')


def set_paragraph_format(paragraph, para_fmt):
    """设置段落级别格式"""
    if not para_fmt:
        return

    # ── 对齐 ──
    alignment = para_fmt.get('alignment')
    if alignment:
        align_map = {
            'left':   WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right':  WD_ALIGN_PARAGRAPH.RIGHT,
            'both':   WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        mapped = align_map.get(alignment.lower())
        if mapped is not None:
            paragraph.alignment = mapped

    pf = paragraph.paragraph_format

    # ── 段前/段后间距（单位: twips, 1/20 pt） ──
    spacing = para_fmt.get('spacing', {})
    if 'before' in spacing:
        pf.space_before = Pt(spacing['before'] / 20.0)
    if 'after' in spacing:
        pf.space_after = Pt(spacing['after'] / 20.0)

    # ── 行距 ──
    if 'line' in spacing:
        line_val = int(spacing['line'])
        line_rule = spacing.get('lineRule', 'auto')
        pf.line_spacing = Pt(line_val / 20.0) if line_rule == 'exact' else float(line_val) / 240.0

    # ── 缩进 ──
    indent = para_fmt.get('indent', {})
    if 'left' in indent:
        pf.left_indent = Twips(indent['left'])
    if 'right' in indent:
        pf.right_indent = Twips(indent['right'])
    if 'firstLine' in indent:
        pf.first_line_indent = Twips(indent['firstLine'])

    # ── 大纲级别 ──
    ol = para_fmt.get('outlineLvl')
    if ol is not None:
        pPr = paragraph._element.get_or_add_pPr()
        el = pPr.find(qn('w:outlineLvl'))
        if el is None:
            el = OxmlElement('w:outlineLvl')
            pPr.append(el)
        el.set(qn('w:val'), str(ol))


def set_cell_shading(cell, fill_color):
    """设置单元格背景底色"""
    if not fill_color or fill_color.lower() == 'auto':
        return
    tcPr = cell._tc.get_or_add_tcPr()
    # 移除旧 shd
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear"/>'
    )
    tcPr.append(shd)


def set_cell_borders(cell, borders):
    """设置单元格边框"""
    if not borders:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, bd in borders.items():
        val = bd.get('val', 'single')
        sz = bd.get('sz', 4)
        space = bd.get('space', 0)
        color = bd.get('color', 'auto')
        b_el = parse_xml(
            f'<{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        )
        tcBorders.append(b_el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, margins):
    """设置单元格边距"""
    if not margins:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
    for key, val in margins.items():
        m_el = parse_xml(f'<{key} {nsdecls("w")} w:w="{val}" w:type="dxa"/>')
        tcMar.append(m_el)
    tcPr.append(tcMar)


# ═══════════════════════════════════════════════════════════
# 主还原函数
# ═══════════════════════════════════════════════════════════

def restore_docx(json_data, output_path):
    """
    从 JSON 数据（dict 或 JSON 文件路径）还原 Word 文档。

    参数:
      json_data : dict | str
        - dict: 已解析的 JSON 对象
        - str:  .json 文件路径
      output_path : str
        输出 .docx 文件路径

    返回: output_path
    """
    # ── 加载 ──
    if isinstance(json_data, str):
        with open(json_data, 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        data = json_data

    # ── 构建文档 ──
    doc = Document()

    # ═══ 页面属性 ═══
    page = data.get('page', {})
    section = doc.sections[0]

    # 页面尺寸（OOXML 单位是 twips，兼容旧字段名 width_emu）
    pg_w = page.get('width_twips') or page.get('width_emu')
    pg_h = page.get('height_twips') or page.get('height_emu')
    if pg_w:
        section.page_width  = Twips(pg_w)
    if pg_h:
        section.page_height = Twips(pg_h)

    margins = page.get('margins', {})
    if margins:
        for key in ('top', 'right', 'bottom', 'left', 'header', 'footer'):
            if key in margins:
                setattr(section, f'{key}_margin', Twips(margins[key]))

    # ═══ 文档级默认字体 ───────────────────────────────
    defaults = data.get('defaults', {})
    default_run = defaults.get('run', {})
    default_font = default_run.get('font', {})
    if default_font:
        # 注入到 Normal 样式
        normal = doc.styles['Normal']
        name = default_font.get('ascii') or default_font.get('eastAsia')
        if name:
            normal.font.name = name
            rPr = normal.element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                normal.element.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            # 先清除主题字体引用
            for theme_attr in ('w:asciiTheme', 'w:eastAsiaTheme', 'w:hAnsiTheme', 'w:cstheme'):
                if rFonts.get(qn(theme_attr)):
                    del rFonts.attrib[qn(theme_attr)]
            for key in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                if default_font.get(key):
                    rFonts.set(qn(f'w:{key}'), default_font[key])
        sz = default_run.get('sz_hp')
        if sz:
            normal.font.size = Pt(sz / 2.0)

    # ═══ 样式注册：将模板 JSON 中的自定义样式应用到文档 ═══
    styles = data.get('styles', {})

    # OOXML 样式名 → python-docx 样式名映射（前者小写，后者首字母大写）
    STYLE_NAME_MAP = {}
    for s in doc.styles:
        STYLE_NAME_MAP[s.name.lower()] = s.name

    for sid, sdef in styles.items():
        sname = sdef.get('name', '')
        stype = sdef.get('type', 'paragraph')
        if not sname or sname == 'Default Paragraph Font':
            continue

        # 查找 python-docx 中的对应样式（大小写不敏感）
        mapped = STYLE_NAME_MAP.get(sname.lower(), sname)
        try:
            doc_style = doc.styles[mapped]
        except KeyError:
            continue

        # ── 字体格式 ──
        run_def = sdef.get('run', {})
        if run_def:
            f = doc_style.font
            sz = run_def.get('sz_hp')
            if sz:
                f.size = Pt(sz / 2.0)
            color = run_def.get('color')
            if color and color != 'auto':
                rgb = hex_to_rgb(color)
                if rgb:
                    f.color.rgb = RGBColor(*rgb)
            if run_def.get('bold') or run_def.get('bold_cjk'):
                f.bold = True
            if run_def.get('italic') or run_def.get('italic_cjk'):
                f.italic = True

        # ── 字体名称（独立处理：先清 theme，有则设显式名，无则保持清除状态继承 Normal）
        rPr = doc_style.element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            doc_style.element.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            # 无论 JSON 有无 font 信息，先清除主题引用（来自 python-docx 默认模板）
            for theme_attr in ('w:asciiTheme', 'w:eastAsiaTheme', 'w:hAnsiTheme', 'w:cstheme'):
                key = qn(theme_attr)
                if key in rFonts.attrib:
                    del rFonts.attrib[key]

        font_map = (run_def or {}).get('font', {})
        if font_map:
            name = font_map.get('ascii') or font_map.get('eastAsia')
            if name:
                f = doc_style.font
                f.name = name
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                if font_map.get('eastAsia'):
                    rFonts.set(qn('w:eastAsia'), font_map['eastAsia'])
                if font_map.get('ascii'):
                    rFonts.set(qn('w:ascii'), font_map['ascii'])
                if font_map.get('hAnsi'):
                    rFonts.set(qn('w:hAnsi'), font_map['hAnsi'])

        # ── 段落格式 ──
        para_def = sdef.get('paragraph', {})
        if para_def and stype == 'paragraph':
            pf = doc_style.paragraph_format
            spacing = para_def.get('spacing', {})
            if 'before' in spacing:
                pf.space_before = Pt(spacing['before'] / 20.0)
            if 'after' in spacing:
                pf.space_after = Pt(spacing['after'] / 20.0)
            ol = para_def.get('outlineLvl')
            if ol is not None:
                pPr = doc_style.element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    doc_style.element.append(pPr)
                ol_el = pPr.find(qn('w:outlineLvl'))
                if ol_el is None:
                    ol_el = OxmlElement('w:outlineLvl')
                    pPr.append(ol_el)
                ol_el.set(qn('w:val'), str(ol))

    # ═══ 处理每个 block ═══
    blocks = data.get('blocks', [])

    for block in blocks:
        block_type = block.get('type')

        # ── 段落 ──────────────────────────────────────
        if block_type == 'paragraph':
            style_id = block.get('style')
            style_name = None
            if style_id and style_id in styles:
                raw_name = styles[style_id].get('name', '')
                style_name = STYLE_NAME_MAP.get(raw_name.lower(), raw_name)

            # 创建段落
            try:
                if style_name:
                    paragraph = doc.add_paragraph(style=style_name)
                else:
                    paragraph = doc.add_paragraph()
            except Exception:
                paragraph = doc.add_paragraph()

            # ── 段落格式 ──
            para_fmt = block.get('paraFmt')
            set_paragraph_format(paragraph, para_fmt)

            # ── 项目编号 (numbering) ──
            if para_fmt and para_fmt.get('numbering'):
                num = para_fmt['numbering']
                pPr = paragraph._element.get_or_add_pPr()
                numPr = OxmlElement('w:numPr')
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'), str(num.get('level', 0)))
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'), str(num.get('numId', '1')))
                numPr.append(ilvl)
                numPr.append(numId)
                pPr.append(numPr)

            # ── Runs ──
            runs = block.get('runs', [])
            for run_data in runs:
                text = run_data.get('text', '')
                fmt  = run_data.get('fmt', {})
                run = paragraph.add_run(text)
                set_run_font(run, fmt)

            # ── 图片嵌入 ──
            image_data = block.get('image')
            image_path = None

            if isinstance(image_data, str):
                image_path = image_data
                width_emu = None
                height_emu = None
            elif isinstance(image_data, dict):
                image_path = image_data.get('path') or image_data.get('src')
                # 兼容只有 rId 占位的情况（跳过）
                if image_path and not isinstance(image_path, str):
                    image_path = None
                width_emu  = image_data.get('width_emu')
                height_emu = image_data.get('height_emu')

            if image_path and os.path.isfile(image_path):
                try:
                    run = paragraph.add_run()
                    if width_emu and height_emu:
                        run.add_picture(image_path,
                                        width=Emu(width_emu),
                                        height=Emu(height_emu))
                    else:
                        run.add_picture(image_path, width=Inches(5.5))
                except Exception as e:
                    print(f"  ⚠ 插入图片失败 [{os.path.basename(image_path)}]: {e}")

        # ── 表格 ──────────────────────────────────────
        elif block_type == 'table':
            num_rows = block.get('numRows', 0)
            num_cols = block.get('numCols', 0)
            columns  = block.get('columns', [])
            rows     = block.get('rows', [])

            # 创建表格
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'

            # ── 表格级格式 ──
            tbl_fmt = block.get('tableFmt', {})
            # 获取或创建 w:tblPr
            tbl_elem = table._tbl
            tblPr = tbl_elem.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl_elem.insert(0, tblPr)  # tblPr 必须是第一个子元素

            # 表格边框
            tbl_borders = tbl_fmt.get('borders', {})
            if tbl_borders:
                borders_el = parse_xml(f'<w:tblBorders {nsdecls("w")}/>')
                for edge, bd in tbl_borders.items():
                    b_el = parse_xml(
                        f'<{edge} {nsdecls("w")} w:val="{bd.get("val","single")}" '
                        f'w:sz="{bd.get("sz",4)}" w:space="{bd.get("space",0)}" '
                        f'w:color="{bd.get("color","auto")}"/>'
                    )
                    borders_el.append(b_el)
                tblPr.append(borders_el)

            # 表格边距
            tbl_margins = tbl_fmt.get('cellMargins', {})
            if tbl_margins:
                mar_el = parse_xml(f'<w:tblCellMar {nsdecls("w")}/>')
                for key, val in tbl_margins.items():
                    m_el = parse_xml(f'<{key} {nsdecls("w")} w:w="{val}" w:type="dxa"/>')
                    mar_el.append(m_el)
                tblPr.append(mar_el)

            # 表格宽度
            if 'width_dxa' in tbl_fmt:
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), str(tbl_fmt['width_dxa']))
                tblW.set(qn('w:type'), tbl_fmt.get('width_type', 'dxa'))
                tblPr.append(tblW)

            # 列宽
            for i, col in enumerate(columns):
                if 'width_dxa' in col:
                    for cell in table.columns[i].cells:
                        cell.width = Twips(col['width_dxa'])

            # ── 填充单元格 ──
            for row_idx, row_data in enumerate(rows):
                if row_idx >= len(table.rows):
                    break
                cells = row_data.get('cells', [])
                for col_idx, cell_data in enumerate(cells):
                    if col_idx >= len(table.columns):
                        break
                    cell = table.cell(row_idx, col_idx)

                    # 彻底清空单元格：删除所有旧段落，重建
                    for old_p in cell.paragraphs:
                        p_elem = old_p._element
                        p_elem.getparent().remove(p_elem)

                    paragraphs = cell_data.get('paragraphs', [])
                    for p_idx, p_data in enumerate(paragraphs):
                        p = cell.add_paragraph()

                        # 段落格式
                        p_fmt = p_data.get('paraFmt')
                        if p_fmt:
                            set_paragraph_format(p, p_fmt)

                        # Runs
                        for run_data in p_data.get('runs', []):
                            text = run_data.get('text', '')
                            fmt  = run_data.get('fmt', {})
                            run = p.add_run(text)
                            set_run_font(run, fmt)

                    # ── 单元格格式 ──
                    cell_fmt = cell_data.get('cellFmt', {})
                    if 'fill' in cell_fmt:
                        set_cell_shading(cell, cell_fmt['fill'])
                    if 'borders' in cell_fmt:
                        set_cell_borders(cell, cell_fmt['borders'])
                    if 'margins' in cell_fmt:
                        set_cell_margins(cell, cell_fmt['margins'])

    # ═══ 保存 ═══
    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════

if __name__ == '__main__':
    import argparse
    p = argparse.ArgumentParser(description='将填报完成的 JSON 还原为 .docx')
    p.add_argument('json', help='JSON 文件路径')
    p.add_argument('--out', '-o', default=None, help='输出 .docx 路径（默认: 同目录下 filled_output.docx）')
    args = p.parse_args()

    src = args.json
    out = args.out or os.path.join(os.path.dirname(os.path.abspath(src)), 'filled_output.docx')

    result = restore_docx(src, out)
    print(f'✅ 文档已生成: {result}')
